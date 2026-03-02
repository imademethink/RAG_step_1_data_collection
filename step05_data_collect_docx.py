from docx import Document
from docx.shared import Inches

# tell top 10 operations to be done on Word file using Python

# 1. Create and Save Documents.
# 2. Add Headings
# 3. Insert Paragraphs
# 4. Text Styling (Runs)
# 5. Add Tables
# 6. Insert Images
# 7. Extract Text
# 8. Find and Replace
# 9. Add Lists
# 10. Page Breaks and Sections:




# 1. Create and Save Documents
doc1 = Document()
doc1.save('my_new_file.docx')


# 2. Add Headings
doc2 = Document()
doc2.add_heading('Main Title', 0)
doc2.add_heading('Section 1', level=1)



# 3. Insert Paragraphs
doc3 = Document()
para = doc3.add_paragraph('This is a simple paragraph.')


# 4. Text Styling (Runs)
doc4 = Document()
para = doc4.add_paragraph('Normal text and ')
run = para.add_run('bold text')
run.bold = True
para.add_run(' and ').add_run('italic text').italic = True


# 5. Add Tables
doc5 = Document()
table = doc5.add_table(rows=2, cols=2)
table.cell(0, 0).text = 'Header 1'
table.cell(0, 1).text = 'Header 2'
table.rows[1].cells[0].text = 'Data A'



# 6. Insert Images
doc6 = Document()
doc6.add_picture('image.png', width=Inches(2.0))


# 7. Extract Text
old_doc = Document('existing.docx')
full_text = [p.text for p in old_doc.paragraphs]
print('\n'.join(full_text))


# 8. Find and Replace
doc7 = Document('existing.docx')
for p in doc7.paragraphs:
    if 'TARGET_WORD' in p.text:
        p.text = p.text.replace('TARGET_WORD', 'REPLACEMENT')


# 9. Add Lists
doc8 = Document('existing.docx')
doc8.add_paragraph('First bullet point', style='List Bullet')
doc8.add_paragraph('First numbered item', style='List Number')


# 10. Page Breaks and Sections:
doc9 = Document('existing.docx')
doc9.add_page_break()

